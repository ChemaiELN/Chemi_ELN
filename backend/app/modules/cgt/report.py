"""Generate a .docx experiment report for an approved CGT experiment.

Reuses ADC's generic docx styling helpers (app.modules.experiments.report)
but iterates CGT's own template shape — section/screen/field keyed by `id`/
`name` (not ADC's `key`), field types are the CGT template-builder's UPPERCASE
enum (SINGLE_LINE_TEXT, NUMBER, DATE, CHECKLIST, ...), and table screens are
signalled by a "(table)"/"(entry table)" title suffix rather than a field type.
"""
import re
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.modules.experiments.report import (
    NAVY, WHITE, SLATE, DARK,
    _shade_cell, _bold_cell, _size_cell, _set_col_width, _add_top_border, _fmt_dt,
)

_TABLE_TITLE_RE = re.compile(r"\((?:entry\s+)?table\)\s*$", re.IGNORECASE)
_ENTRY_TABLE_TITLE_RE = re.compile(r"\(entry\s+table\)\s*$", re.IGNORECASE)


def _clean_title(title: str) -> str:
    return _TABLE_TITLE_RE.sub("", title or "").strip()


def _fmt_field(value: Any, ftype: str) -> str:
    if value is None or value == "":
        return "—"
    if ftype == "CHECKBOX":
        return "Yes" if value else "No"
    if ftype == "CHECKLIST":
        return ", ".join(value) if isinstance(value, list) and value else "—"
    if ftype in ("DATE", "DATE_TIME") and isinstance(value, str):
        return _fmt_dt(value)
    if isinstance(value, list):
        return f"({len(value)} row(s))"
    return str(value)


def generate_cgt_experiment_docx(
    experiment: dict,
    notebook: dict,
    project: dict,
    approver_name: str | None,
    submitter_name: str | None,
) -> bytes:
    import io

    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.8)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Cover ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LAURUS ELN  —  CGT MODULE")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.all_caps = True
    _add_top_border(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Experiment Report")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{experiment.get('full_code', '')}  ·  {experiment.get('title', '')}")
    r.font.size = Pt(12)
    r.font.color.rgb = SLATE

    doc.add_paragraph()

    # ── Metadata table ──────────────────────────────────────────────────────
    meta_pairs = [
        ("Project",         f"{project.get('code', '')} — {project.get('name', '')}"),
        ("Notebook",        f"{notebook.get('code', '')} — {notebook.get('title', '')}"),
        ("Experiment Code", experiment.get("full_code", "")),
        ("Status",          experiment.get("status", "")),
        ("Submitted By",    submitter_name or "—"),
        ("Submitted At",    _fmt_dt(experiment.get("submitted_at"))),
        ("Approved By",     approver_name or "—"),
        ("Approved At",     _fmt_dt(experiment.get("approved_at"))),
        ("Generated",       datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
    ]
    mt = doc.add_table(rows=0, cols=2)
    mt.style = "Table Grid"
    for label, value in meta_pairs:
        row = mt.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        _bold_cell(row.cells[0], size_pt=9, color=NAVY)
        _size_cell(row.cells[1], size_pt=9)
        _shade_cell(row.cells[0], "EFF6FF")
    _set_col_width(mt, 0, 1.8)
    _set_col_width(mt, 1, 4.5)
    doc.add_paragraph()

    for heading_text, value_key in [("Observations", "observations"), ("Conclusion", "conclusion")]:
        val = experiment.get(value_key)
        if val:
            h = doc.add_heading(heading_text, level=2)
            if h.runs:
                h.runs[0].font.color.rgb = NAVY
            doc.add_paragraph(str(val))

    # ── Section data ────────────────────────────────────────────────────────
    snapshot = notebook.get("template_snapshot") or {}
    sections = snapshot.get("sections", [])
    exp_data: dict = experiment.get("data") or {}

    for sec_def in sections:
        section_id = sec_def.get("id", "")
        section_data = exp_data.get(section_id) or {}
        if not section_data:
            continue  # skip sections with no data entered for this experiment

        sec_title = sec_def.get("title") or "Section"
        sh = doc.add_heading(sec_title, level=1)
        if sh.runs:
            sh.runs[0].font.color.rgb = NAVY

        for screen in sec_def.get("screens", []):
            screen_id  = screen.get("id", "")
            raw_title  = screen.get("title") or "Screen"
            sdata      = section_data.get(screen_id)
            fields     = [f for f in screen.get("fields", []) if f.get("type") not in ("SECTION_HEADING", "SPACER")]
            is_table   = bool(_TABLE_TITLE_RE.search(raw_title))

            if sdata is None or (isinstance(sdata, list) and not sdata) or (isinstance(sdata, dict) and not sdata):
                continue

            h = doc.add_heading(_clean_title(raw_title), level=2)
            if h.runs:
                h.runs[0].font.color.rgb = NAVY

            if is_table:
                rows = sdata if isinstance(sdata, list) else []
                if fields and rows:
                    col_count = len(fields)
                    tbl = doc.add_table(rows=1, cols=col_count)
                    tbl.style = "Table Grid"
                    hcells = tbl.rows[0].cells
                    for ci, f in enumerate(fields):
                        hcells[ci].text = f.get("label") or f.get("name", "")
                        _shade_cell(hcells[ci], "1E3A5F")
                        _bold_cell(hcells[ci], size_pt=8, color=WHITE)
                    for ri, row_data in enumerate(rows):
                        dr_cells = tbl.add_row().cells
                        for ci, f in enumerate(fields):
                            fval = _fmt_field(row_data.get(f.get("name", "")), f.get("type", ""))
                            dr_cells[ci].text = fval
                            _size_cell(dr_cells[ci], size_pt=8.5)
                            if ri % 2 == 1:
                                _shade_cell(dr_cells[ci], "F8FAFC")
                    avail = 6.3 / col_count
                    for ci in range(col_count):
                        _set_col_width(tbl, ci, avail)
                doc.add_paragraph()
            else:
                values = sdata if isinstance(sdata, dict) else {}
                if fields:
                    ft = doc.add_table(rows=1, cols=2)
                    ft.style = "Table Grid"
                    hcells = ft.rows[0].cells
                    hcells[0].text = "Field"
                    hcells[1].text = "Value"
                    for hc in hcells:
                        _shade_cell(hc, "1E3A5F")
                        _bold_cell(hc, size_pt=8.5, color=WHITE)
                    for f in fields:
                        flabel = f.get("label") or f.get("name", "")
                        fval = _fmt_field(values.get(f.get("name", "")), f.get("type", ""))
                        dr = ft.add_row()
                        dr.cells[0].text = flabel
                        dr.cells[1].text = fval
                        _bold_cell(dr.cells[0], size_pt=9)
                        _shade_cell(dr.cells[0], "F1F5F9")
                        _size_cell(dr.cells[1], size_pt=9)
                    _set_col_width(ft, 0, 2.2)
                    _set_col_width(ft, 1, 4.1)
                    doc.add_paragraph()

    doc.add_paragraph()
    fp = doc.add_paragraph()
    _add_top_border(fp)
    fr = fp.add_run(f"Generated by Laurus ELN  ·  {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}  ·  Confidential")
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = SLATE

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
