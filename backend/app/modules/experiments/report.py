"""Generate a .docx experiment report for an approved experiment."""
import io
import re
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ──────────────────────────────────────────────────────────────────

def _strip_html(html: str) -> str:
    if not html:
        return ""
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    text = re.sub(r"</p>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<li[^>]*>", "• ", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _fmt(val: Any, ftype: str = "text") -> str:
    if val is None or val == "":
        return "—"
    if ftype == "boolean":
        return "Yes" if val else "No"
    if ftype in ("rich_text", "textarea"):
        return _strip_html(str(val))
    if isinstance(val, list):
        return f"({len(val)} row(s))"
    if isinstance(val, dict):
        import json
        return json.dumps(val)
    return str(val)


def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_width(table, col_idx: int, width_inches: float):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def _bold_cell(cell, size_pt: float = 9.5, color: RGBColor | None = None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(size_pt)
            if color:
                run.font.color.rgb = color


def _size_cell(cell, size_pt: float = 9.5):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(size_pt)


NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
VIOLET = RGBColor(0x7C, 0x3A, 0xED)
SLATE  = RGBColor(0x47, 0x55, 0x69)
DARK   = RGBColor(0x1E, 0x29, 0x3B)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_experiment_docx(
    experiment: dict,
    notebook: dict,
    project: dict,
    approver_name: str | None,
    submitter_name: str | None,
) -> bytes:
    doc = Document()

    # Page setup — A4
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.8)

    # Base font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Cover ──────────────────────────────────────────────────────────────

    # Top label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LAURUS ELN  —  ADC MODULE")
    r.font.size  = Pt(8)
    r.font.bold  = True
    r.font.color.rgb = VIOLET
    r.font.all_caps  = True

    # Divider line via paragraph border
    _add_top_border(p)

    # Big title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Experiment Report")
    r.font.size  = Pt(24)
    r.font.bold  = True
    r.font.color.rgb = DARK

    # Experiment code · title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_str = f"{experiment.get('full_code', '')}  ·  {experiment.get('title', '')}"
    r = p.add_run(code_str)
    r.font.size  = Pt(12)
    r.font.color.rgb = SLATE

    doc.add_paragraph()  # spacer

    # ── Metadata table ──────────────────────────────────────────────────────

    meta_pairs = [
        ("Project",         f"{project.get('code', '')} — {project.get('name', '')}"),
        ("Notebook",        f"{notebook.get('code', '')} — {notebook.get('title', '')}"),
        ("Section",         (experiment.get("title") or "").replace("_", " ").title() or experiment.get("section_key", "").replace("_", " ").title()),
        ("Experiment Code", experiment.get("full_code", "")),
        ("Status",          experiment.get("status", "")),
        ("Submitted By",    submitter_name or "—"),
        ("Submitted At",    _fmt_dt(experiment.get("submitted_at"))),
        ("Approved By",     approver_name or "—"),
        ("Approved At",     _fmt_dt(experiment.get("approved_at"))),
        ("Generated",       datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
    ]

    mt = doc.add_table(rows=0, cols=2)
    mt.style = "Table Grid"

    for label, value in meta_pairs:
        row = mt.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        _bold_cell(row.cells[0], size_pt=9, color=NAVY)
        _size_cell(row.cells[1], size_pt=9)
        _shade_cell(row.cells[0], "EFF6FF")

    _set_col_width(mt, 0, 1.8)
    _set_col_width(mt, 1, 4.5)

    doc.add_paragraph()

    # Disposition
    disp = experiment.get("disposition")
    if disp:
        p = doc.add_paragraph()
        r = p.add_run("Disposition: ")
        r.font.bold = True
        r.font.color.rgb = NAVY
        p.add_run(disp)

    # Observations & Conclusion
    for heading_text, value_key in [("Observations", "observations"), ("Conclusion", "conclusion")]:
        val = experiment.get(value_key)
        if val:
            h = doc.add_heading(heading_text, level=2)
            h.runs[0].font.color.rgb = NAVY
            doc.add_paragraph(_strip_html(val))

    # ── Section data ────────────────────────────────────────────────────────

    snapshot     = notebook.get("template_snapshot") or {}
    all_sections = snapshot.get("sections", [])
    section_key  = experiment.get("section_key")

    # If section_key is set, render just that section.
    # If None (full-notebook experiment), render ALL sections with all screens.
    if section_key:
        sections_to_render = [s for s in all_sections if s.get("key") == section_key]
    else:
        sections_to_render = all_sections

    exp_data: dict = experiment.get("data") or {}

    for sec_def in sections_to_render:
        # Section-level heading when rendering multiple sections
        if not section_key:
            sec_title = sec_def.get("title") or sec_def.get("key", "").replace("_", " ").title()
            sh = doc.add_heading(sec_title, level=1)
            if sh.runs:
                sh.runs[0].font.color.rgb = NAVY

        for screen in sec_def.get("screens", []):
            skey   = screen.get("key", "")
            stitle = screen.get("title") or skey.replace("_", " ").title()
            sdata  = exp_data.get(skey) or {}
            fields = screen.get("fields", [])

            # Skip screens with no data entered at all
            if not sdata:
                continue

            h = doc.add_heading(stitle, level=2)
            if h.runs:
                h.runs[0].font.color.rgb = NAVY

            simple = [f for f in fields if f.get("type") != "table"]
            tables = [f for f in fields if f.get("type") == "table"]

            # Simple fields → label/value table
            if simple:
                ft = doc.add_table(rows=1, cols=2)
                ft.style = "Table Grid"
                # Header row
                hcells = ft.rows[0].cells
                hcells[0].text = "Field"
                hcells[1].text = "Value"
                for hc in hcells:
                    _shade_cell(hc, "1E3A5F")
                    _bold_cell(hc, size_pt=8.5, color=WHITE)

                for f in simple:
                    fkey   = f.get("key", "")
                    flabel = f.get("label") or fkey.replace("_", " ").title()
                    ftype  = f.get("type", "text")
                    unit   = f.get("unit", "")
                    raw    = sdata.get(fkey)
                    fval   = _fmt(raw, ftype)
                    if unit and fval != "—":
                        fval = f"{fval} {unit}"

                    dr = ft.add_row()
                    dr.cells[0].text = flabel
                    dr.cells[1].text = fval
                    _bold_cell(dr.cells[0], size_pt=9)
                    _shade_cell(dr.cells[0], "F1F5F9")
                    _size_cell(dr.cells[1], size_pt=9)

                _set_col_width(ft, 0, 2.2)
                _set_col_width(ft, 1, 4.1)
                doc.add_paragraph()

            # Table fields → full table
            for tf in tables:
                fkey    = tf.get("key", "")
                flabel  = tf.get("label") or fkey.replace("_", " ").title()
                columns = tf.get("columns", [])
                rows    = sdata.get(fkey) or []
                if not isinstance(rows, list):
                    rows = []

                # Sub-heading for the table
                sp = doc.add_paragraph()
                sr = sp.add_run(flabel)
                sr.font.bold  = True
                sr.font.size  = Pt(10)
                sr.font.color.rgb = NAVY

                if columns and rows:
                    col_count = len(columns)
                    tbl = doc.add_table(rows=1, cols=col_count)
                    tbl.style = "Table Grid"

                    # Header row
                    hcells = tbl.rows[0].cells
                    for ci, col in enumerate(columns):
                        hcells[ci].text = col.get("label") or col.get("key", "").replace("_", " ").title()
                        _shade_cell(hcells[ci], "1E3A5F")
                        _bold_cell(hcells[ci], size_pt=8, color=WHITE)

                    # Data rows
                    for ri, dr_data in enumerate(rows):
                        dr_cells = tbl.add_row().cells
                        for ci, col in enumerate(columns):
                            ckey  = col.get("key", "")
                            ctype = col.get("type", "text")
                            cval  = _fmt(dr_data.get(ckey), ctype)
                            dr_cells[ci].text = cval
                            _size_cell(dr_cells[ci], size_pt=8.5)
                            if ri % 2 == 1:
                                _shade_cell(dr_cells[ci], "F8FAFC")

                    # Even column width
                    avail = 6.3 / col_count
                    for ci in range(col_count):
                        _set_col_width(tbl, ci, avail)

                doc.add_paragraph()

    # Footer note
    doc.add_paragraph()
    fp = doc.add_paragraph()
    _add_top_border(fp)
    fr = fp.add_run(f"Generated by Laurus ELN  ·  {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}  ·  Confidential")
    fr.font.size  = Pt(7.5)
    fr.font.color.rgb = SLATE

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _fmt_dt(dt_str: str | None) -> str:
    if not dt_str:
        return "—"
    try:
        return dt_str[:19].replace("T", " ") + " UTC"
    except Exception:
        return str(dt_str)


def _add_top_border(para):
    """Add a top border line to a paragraph (acts as a horizontal rule)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "C7D2FE")  # indigo-200
    pBdr.append(top)
    pPr.append(pBdr)
